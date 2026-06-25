from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs"
OUT_FILE = OUT_DIR / "BẢN KẾ HOẠCH THIẾT KẾ CHUẨN V1.1 VỀ TOOL.docx"

PRIMARY = RGBColor(37, 99, 235)
ACCENT = RGBColor(249, 115, 22)
GREEN = RGBColor(16, 185, 129)
DARK = RGBColor(17, 24, 39)
MUTED = RGBColor(75, 85, 99)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell(cell, text: str, fill: str = "FFFFFF", bold: bool = False, color=DARK) -> None:
    cell.text = ""
    shade_cell(cell, fill)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for idx, line in enumerate(text.split("\n")):
        if idx:
            paragraph.add_run().add_break()
        run = paragraph.add_run(line)
        run.font.name = "Times New Roman"
        run.font.size = Pt(14)
        run.bold = bold
        run.font.color.rgb = color


def add_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(20)
    r.bold = True
    r.font.color.rgb = PRIMARY


def add_subtitle(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(14)
    r.italic = True
    r.font.color.rgb = MUTED


def add_heading(doc: Document, text: str, color=PRIMARY) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(16)
    r.bold = True
    r.font.color.rgb = color


def add_para(doc: Document, text: str, bold: bool = False, color=DARK) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(14)
    r.bold = bold
    r.font.color.rgb = color


def add_bullet(doc: Document, text: str, color=DARK, bold: bool = False) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(14)
    r.bold = bold
    r.font.color.rgb = color


def add_note(doc: Document, title: str, text: str, fill: str = "FFF7ED") -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell(cell, f"{title}\n{text}", fill=fill, bold=True, color=DARK)


def add_table(doc: Document, headers, rows, header_fill="EFF6FF") -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for col, header in enumerate(headers):
        set_cell(table.cell(0, col), header, fill=header_fill, bold=True, color=PRIMARY)
    for row in rows:
        cells = table.add_row().cells
        for col, value in enumerate(row):
            set_cell(cells[col], str(value), fill="FFFFFF", bold=False, color=DARK)


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2)

    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(14)

    add_title(doc, "BẢN KẾ HOẠCH THIẾT KẾ CHUẨN V1.1 VỀ TOOL")
    add_subtitle(doc, "EduPlan AI - Web app tạo kế hoạch bài dạy theo Công văn 2345")
    add_note(
        doc,
        "ĐIỂM CẦN CHÚ Ý",
        "Bản v1.1 đã chỉnh theo hướng: form nhập gọn hơn, nội dung bài học lấy từ OCR ảnh/PDF, Gemini API key do user tự nhập, OpenAI API key do chủ app cấu hình cố định ở server.",
        "FEF3C7",
    )

    add_heading(doc, "1. Mục Tiêu Sản Phẩm")
    for item in [
        "Xây web app giúp giáo viên tạo kế hoạch bài dạy theo Công văn 2345.",
        "Giáo viên chỉ cần nhập thông tin cơ bản và upload ảnh/PDF SGK.",
        "AI tự OCR nội dung bài học, tự suy luận yêu cầu cần đạt nếu tài liệu không có sẵn.",
        "Giáo án sinh ra phải đẹp trên web preview và đẹp khi xuất DOCX/PDF.",
        "Hỗ trợ tinh chỉnh nhanh: rút gọn, chi tiết hơn, thi giảng, thêm năng lực số.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "2. Hướng Triển Khai")
    for item in [
        "Làm bản Web trước, chưa làm EXE/Android ở MVP.",
        "Ưu tiên desktop/laptop, có responsive cho mobile.",
        "Sau khi web ổn định mới tính PWA, desktop hoặc mobile.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "3. Công Nghệ Đề Xuất")
    add_table(
        doc,
        ["Hạng mục", "Công nghệ"],
        [
            ["Frontend", "Next.js + TypeScript"],
            ["UI", "Tailwind CSS + shadcn/ui"],
            ["Backend", "Next.js API routes"],
            ["OCR", "Gemini 2.5 Flash Vision, user tự nhập Gemini API key"],
            ["Sinh giáo án", "OpenAI API cố định ở server"],
            ["Export Word", "docx"],
            ["Export PDF", "Puppeteer render HTML/CSS"],
            ["Lưu nháp/version", "localStorage hoặc IndexedDB"],
        ],
    )

    add_heading(doc, "4. Phạm Vi MVP")
    for item in [
        "Form nhập liệu gọn.",
        "Upload ảnh SGK/PDF.",
        "OCR bằng Gemini key của user.",
        "Preview text OCR để user kiểm tra/sửa.",
        "AI sinh giáo án chuẩn CV2345, đủ 5 phần.",
        "Bảng hoạt động đúng 2 cột: Hoạt động của giáo viên và Hoạt động của học sinh.",
        "Preview đẹp có màu sắc, tiêu đề nổi bật, bảng rõ ràng.",
        "Xuất Word/PDF giữ định dạng đẹp.",
        "Autosave form và version giáo án.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "5. LessonForm Sau Khi Chỉnh")
    add_note(
        doc,
        "THAY ĐỔI QUAN TRỌNG",
        "Không đưa Thời lượng, Nội dung bài học, Yêu cầu cần đạt vào form chính. Thời lượng mặc định 35 phút/tiết. Nội dung và yêu cầu cần đạt lấy từ OCR hoặc để AI tự suy luận.",
        "ECFDF5",
    )
    add_table(
        doc,
        ["Trường", "Cách xử lý"],
        [
            ["Môn học", "Bắt buộc nhập"],
            ["Lớp", "Bắt buộc nhập"],
            ["Tên bài", "Bắt buộc nhập"],
            ["Bộ sách", "Khuyến nghị nhập/chọn"],
            ["Số tiết", "Bắt buộc nhập"],
            ["Ảnh SGK/PDF bài học", "Upload để OCR"],
            ["Đối tượng học sinh", "Chọn mẫu hoặc Auto"],
            ["Môi trường học", "Chọn mẫu hoặc Auto"],
            ["Cơ sở vật chất", "Chọn mẫu/tick nhiều lựa chọn hoặc Auto"],
            ["Phong cách giáo án", "Chọn mẫu và có ô nhập thêm"],
            ["Yêu cầu đặc biệt", "Không bắt buộc"],
            ["Gemini API key cho OCR", "User tự nhập"],
        ],
    )

    add_heading(doc, "6. Giá Trị Gợi Ý Trong Form")
    add_table(
        doc,
        ["Nhóm", "Gợi ý"],
        [
            ["Đối tượng học sinh", "Auto; học sinh vùng núi; nông thôn; thành thị; học lực không đồng đều; khá giỏi; cần hỗ trợ nhiều; khác"],
            ["Môi trường học", "Auto; trường thành phố; trường nông thôn; trường vùng núi; điểm trường lẻ; lớp đông; lớp ít học sinh; khác"],
            ["Cơ sở vật chất", "Auto; TV; Wifi; máy chiếu; bảng tương tác; máy tính giáo viên; loa; phiếu học tập; không có thiết bị trình chiếu; khác"],
            ["Phong cách giáo án", "Chuẩn nộp giảng viên; dạy thật trên lớp; thi giảng/dự giờ; ngắn gọn thực tế; chi tiết nhiều hoạt động; sáng tạo sinh động; khác"],
        ],
    )
    add_bullet(doc, "Checkbox nên có: Cho phép AI tự suy luận phần còn thiếu từ ảnh SGK và thông tin bài học.", ACCENT, True)
    add_bullet(doc, "Checkbox này mặc định bật để giảm gánh nặng nhập liệu cho giáo viên.", ACCENT, True)

    add_heading(doc, "7. Luồng Xử Lý Chính")
    flow = [
        "User nhập thông tin cơ bản.",
        "User nhập Gemini API key.",
        "User upload ảnh/PDF SGK.",
        "Hệ thống validate dữ liệu bắt buộc.",
        "Hệ thống gửi ảnh/PDF đến Gemini OCR bằng key user.",
        "Hệ thống hiển thị text OCR để user rà soát/sửa.",
        "GPT nano làm sạch OCR nếu cần.",
        "GPT mini sinh giáo án JSON theo CV2345.",
        "Quality check kiểm tra đủ phần, đúng bảng, không quá chung chung.",
        "Render preview đẹp.",
        "User chỉnh sửa/tinh chỉnh bằng toolbar.",
        "User xuất Word/PDF hoặc copy nội dung.",
    ]
    for idx, item in enumerate(flow, start=1):
        add_bullet(doc, f"Bước {idx}: {item}")

    add_heading(doc, "8. Model Routing Policy")
    add_table(
        doc,
        ["Tính năng", "Model/API", "Ghi chú"],
        [
            ["OCR ảnh/PDF", "Gemini 2.5 Flash Vision", "Dùng key của user"],
            ["Làm sạch OCR", "GPT nano", "Không thêm ý mới"],
            ["Soạn giáo án thường", "GPT mini", "Model chính"],
            ["Rút gọn/viết lại nhẹ", "GPT nano", "Tiết kiệm chi phí"],
            ["Chi tiết hơn", "GPT mini", "Thêm câu hỏi gợi mở, sản phẩm học tập"],
            ["Thêm năng lực số", "GPT mini", "Chèn phù hợp, không gượng ép"],
            ["Thi giảng/demo", "GPT full", "Chỉ dùng khi user chọn"],
            ["Fallback chất lượng cao", "GPT full", "Khi quality check fail nhiều lần"],
            ["Quality check", "Rule-based", "Chỉ gọi AI khi cần sửa cấu trúc"],
        ],
    )

    add_heading(doc, "9. Chi Phí Dự Kiến")
    add_note(
        doc,
        "LƯU Ý CHI PHÍ",
        "OCR dùng Gemini key của user nên chi phí chính chủ app chịu là OpenAI. Các số dưới đây là ước tính thiết kế, có thể thay đổi theo giá API và độ dài nội dung.",
        "FFF7ED",
    )
    add_table(
        doc,
        ["Loại giáo án", "Chi phí dự kiến"],
        [
            ["Giáo án thường, nhập text hoặc OCR đã có", "Khoảng 75-375 VNĐ/lần"],
            ["Giáo án thường, có làm sạch OCR", "Khoảng 125-550 VNĐ/lần"],
            ["Chi tiết hơn/thêm năng lực số", "Khoảng 150-700 VNĐ/lần"],
            ["Thi giảng/demo cao cấp", "Khoảng 1.000-4.250 VNĐ/lần"],
        ],
    )

    add_heading(doc, "10. Thiết Kế Giao Diện")
    for item in [
        "Layout desktop 2 cột: trái là form nhập liệu, phải là preview giáo án.",
        "Mobile dùng tab: Nhập liệu, Xem trước, Xuất file.",
        "Theme sáng học thuật, tiêu đề phần có màu nhấn.",
        "Khối quan trọng có nền màu nhạt.",
        "Bảng 2 cột GV/HS có header màu, border sạch.",
        "Toolbar preview gồm: Tạo lại hay hơn, Rút gọn, Chi tiết hơn, Thêm năng lực số, Thi giảng, Xuất Word, Xuất PDF, Copy toàn bộ.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "11. Thiết Kế Giáo Án Đẹp")
    for item in [
        "Preview không phải text thô, mà render theo template có màu sắc.",
        "Tiêu đề bài nổi bật.",
        "Các phần I-V rõ ràng.",
        "Màu nhấn xanh/cam nhẹ.",
        "Bảng hoạt động đẹp, dễ đọc.",
        "Vùng Yêu cầu cần đạt và Đánh giá có nền màu nhẹ.",
        "Có thể chèn ảnh user upload vào cột GV nếu phù hợp nội dung hoạt động.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "12. Xuất Word/PDF")
    add_table(
        doc,
        ["Định dạng", "Yêu cầu"],
        [
            ["Word", "Times New Roman, cỡ chữ body 13, khoảng cách đoạn 3pt, tiêu đề màu, bảng 2 cột rõ, khổ A4, footer nhỏ"],
            ["PDF", "Render gần giống preview web, giữ màu sắc, không lỗi font tiếng Việt, không vỡ bảng, dễ in"],
        ],
    )

    add_heading(doc, "13. API Nội Bộ")
    for item in [
        "POST /api/ocr",
        "POST /api/lesson/generate",
        "POST /api/lesson/refine",
        "POST /api/quality-check",
        "POST /api/export/docx",
        "POST /api/export/pdf",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "14. Schema Đầu Vào Rút Gọn")
    add_para(doc, "LessonInput gồm: subject, grade, lessonTitle, book, periods, duration mặc định 35, ocrText, studentProfile, teachingEnvironment, facilities, style, customStyle, specialRequest, allowAiInference, geminiApiKey.")

    add_heading(doc, "15. Schema Đầu Ra")
    add_para(doc, "LessonPlan gồm: generalInfo, outcomes, materials, activities, assessment, adjustments, contextFit, meta.")
    add_para(doc, "activities phải có phase, title, objective, teacherActions, studentActions và learningProducts nếu có.", True, ACCENT)

    add_heading(doc, "16. Quality Guardrails")
    for item in [
        "Có đủ 5 phần I-V.",
        "Có bảng hoạt động 2 cột.",
        "Có đủ hoạt động: khởi động, khám phá, luyện tập, vận dụng, đánh giá.",
        "Có đánh giá học sinh.",
        "Có điều chỉnh sau bài dạy.",
        "Có đồ dùng giáo viên/học sinh.",
        "Có yếu tố phù hợp hoàn cảnh dạy học.",
        "Không viết quá chung chung.",
        "Không cho xuất nếu thiếu cấu trúc nghiêm trọng.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "17. Bảo Mật")
    for item in [
        "OpenAI API key đặt trong .env server.",
        "Gemini API key của user không lưu database ở MVP.",
        "Gemini key có thể lưu local ở trình duyệt nếu user chọn ghi nhớ.",
        "Backend dùng Gemini key để OCR rồi không lưu lại.",
        "Có cảnh báo rõ khi user nhập Gemini key.",
        "Có rate limit API generate/refine và giới hạn dung lượng upload.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "18. Kế Hoạch Code Theo Mốc")
    add_table(
        doc,
        ["Mốc", "Thời gian", "Nội dung"],
        [
            ["Mốc 1", "2-3 ngày", "Khởi tạo Next.js, dựng layout, LessonForm rút gọn, LessonPreview mẫu, autosave, validate, mock generate"],
            ["Mốc 2", "5-7 ngày", "API OCR Gemini key user, OCR review, clean OCR, generate bằng OpenAI, render JSON thật, quality check cơ bản"],
            ["Mốc 3", "10-14 ngày", "Export DOCX/PDF, refine toolbar, version history, fix layout/font/bảng, test 10 case mẫu"],
        ],
    )

    add_heading(doc, "19. Tiêu Chí Hoàn Thành MVP")
    for item in [
        "User nhập form ngắn gọn và upload ảnh/PDF.",
        "OCR được bằng Gemini key user.",
        "AI sinh giáo án đúng CV2345.",
        "Nếu thiếu yêu cầu cần đạt, AI tự suy luận được.",
        "Preview đẹp, có màu sắc.",
        "Xuất Word/PDF đẹp, không lỗi font.",
        "Bảng 2 cột GV/HS chuẩn.",
        "Có refine cơ bản, autosave và version.",
        "Thời gian tạo giáo án trung bình dưới 60 giây.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "20. Kết Luận", GREEN)
    add_note(
        doc,
        "KẾT LUẬN",
        "Kế hoạch v1.1 tập trung vào trải nghiệm giáo viên: nhập ít hơn, upload ảnh nhiều hơn, AI tự suy luận phần còn thiếu, giáo án đẹp và dùng được. OCR dùng Gemini key của user, OpenAI được quản lý an toàn ở backend.",
        "ECFDF5",
    )

    doc.save(OUT_FILE)
    print("created plan v1.1 docx")


if __name__ == "__main__":
    main()
