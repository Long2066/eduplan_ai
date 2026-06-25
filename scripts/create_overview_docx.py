from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs"
OUT_FILE = OUT_DIR / "EDUPLAN_AI_SO_DO_TONG_QUAN.docx"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, fill: str = "FFFFFF", size: int = 11, bold: bool = True) -> None:
    cell.text = ""
    set_cell_shading(cell, fill)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for index, line in enumerate(text.split("\n")):
        if index:
            paragraph.add_run().add_break()
        run = paragraph.add_run(line)
        run.font.name = "Times New Roman"
        run.font.size = Pt(size)
        run.bold = bold
        run.font.color.rgb = RGBColor(17, 24, 39)


def add_heading(doc: Document, text: str, size: int = 16) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(37, 99, 235)


def add_paragraph(doc: Document, text: str, size: int = 12, bold: bool = False, center: bool = False) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor(55, 65, 81)


def style_table(table) -> None:
    table.autofit = True
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(3)


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)

    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Cm(1.4)
    section.bottom_margin = Cm(1.4)
    section.left_margin = Cm(1.4)
    section.right_margin = Cm(1.4)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("SƠ ĐỒ TỔNG QUAN CÔNG CỤ EDUPLAN AI")
    title_run.bold = True
    title_run.font.name = "Times New Roman"
    title_run.font.size = Pt(22)
    title_run.font.color.rgb = RGBColor(37, 99, 235)

    add_paragraph(
        doc,
        "Web app tạo kế hoạch bài dạy theo Công văn 2345, có OCR ảnh SGK/PDF, AI sinh giáo án, kiểm tra chất lượng và xuất Word/PDF đẹp.",
        12,
        center=True,
    )

    add_heading(doc, "1. Sơ Đồ Cấu Trúc Tổng Quan")
    add_paragraph(
        doc,
        "Nhìn từ trái sang phải: giáo viên nhập dữ liệu, web app xử lý giao diện, backend gọi AI/OCR, kiểm tra chất lượng, sau đó xuất file.",
        11,
    )

    table = doc.add_table(rows=3, cols=7)
    table.style = "Table Grid"
    style_table(table)

    set_cell_text(table.cell(0, 0), "GIÁO VIÊN\nNhập bài học\nUpload SGK/PDF", "FEF3C7")
    set_cell_text(table.cell(0, 1), "→", "FFFFFF", 18)
    set_cell_text(table.cell(0, 2), "WEB APP\nForm nhập liệu\nPreview đẹp\nToolbar tinh chỉnh", "EFF6FF")
    set_cell_text(table.cell(0, 3), "→", "FFFFFF", 18)
    set_cell_text(table.cell(0, 4), "BACKEND API\nValidate\nGenerate\nRefine\nExport", "F3E8FF")
    set_cell_text(table.cell(0, 5), "→", "FFFFFF", 18)
    set_cell_text(table.cell(0, 6), "AI / OCR\nGemini Vision\nGPT nano / mini / full", "ECFDF5")

    set_cell_text(table.cell(1, 0), "", "FFFFFF", 10, False)
    set_cell_text(table.cell(1, 1), "", "FFFFFF", 10, False)
    set_cell_text(table.cell(1, 2), "↓", "FFFFFF", 18)
    set_cell_text(table.cell(1, 3), "", "FFFFFF", 10, False)
    set_cell_text(table.cell(1, 4), "↓", "FFFFFF", 18)
    set_cell_text(table.cell(1, 5), "", "FFFFFF", 10, False)
    set_cell_text(table.cell(1, 6), "↓", "FFFFFF", 18)

    set_cell_text(table.cell(2, 0), "LƯU NHÁP\nAutosave\nVersion history", "F9FAFB")
    set_cell_text(table.cell(2, 1), "←", "FFFFFF", 18)
    set_cell_text(table.cell(2, 2), "PREVIEW GIÁO ÁN\nMàu sắc đẹp\nBảng 2 cột GV/HS", "EFF6FF")
    set_cell_text(table.cell(2, 3), "←", "FFFFFF", 18)
    set_cell_text(table.cell(2, 4), "QUALITY CHECK\nĐủ 5 phần\nĐúng 2 cột\nCó đánh giá", "ECFDF5")
    set_cell_text(table.cell(2, 5), "→", "FFFFFF", 18)
    set_cell_text(table.cell(2, 6), "ĐẦU RA\nWord .docx\nPDF\nCopy nội dung", "FFF7ED")

    add_heading(doc, "2. Luồng Hoạt Động Chính")
    flow = doc.add_table(rows=8, cols=3)
    flow.style = "Table Grid"
    style_table(flow)

    headers = ["Bước", "Khối xử lý", "Mô tả dễ hiểu"]
    for col, header in enumerate(headers):
        set_cell_text(flow.cell(0, col), header, "EFF6FF", 12)

    rows = [
        ("1", "Nhập liệu", "Giáo viên nhập môn, lớp, tên bài, số tiết, yêu cầu cần đạt, bối cảnh dạy học và có thể upload ảnh SGK/PDF.", "FEF3C7"),
        ("2", "OCR", "Gemini Vision đọc ảnh/PDF thành văn bản. Người dùng được xem và sửa text OCR trước khi tạo giáo án.", "ECFDF5"),
        ("3", "Làm sạch text", "GPT nano sửa lỗi OCR, chuẩn hóa dấu câu, xuống dòng, nhưng không tự thêm nội dung mới.", "ECFDF5"),
        ("4", "Sinh giáo án", "GPT mini sinh giáo án JSON theo CV2345. GPT full chỉ dùng cho thi giảng/demo hoặc khi chất lượng chưa đạt.", "ECFDF5"),
        ("5", "Kiểm tra chất lượng", "Hệ thống kiểm tra đủ 5 phần, đúng bảng 2 cột GV/HS, có đánh giá, có điều chỉnh và phù hợp hoàn cảnh giảng dạy.", "ECFDF5"),
        ("6", "Preview & tinh chỉnh", "Giáo án hiển thị đẹp trên web. Người dùng có thể rút gọn, chi tiết hơn, thêm năng lực số, chuyển sang thi giảng.", "EFF6FF"),
        ("7", "Xuất file", "Xuất Word/PDF với Times New Roman 13, tiêu đề màu, bảng 2 cột rõ, khổ A4, không lỗi font tiếng Việt.", "FFF7ED"),
    ]

    for row_index, (step, block, desc, fill) in enumerate(rows, start=1):
        set_cell_text(flow.cell(row_index, 0), step, "FFFFFF", 12)
        set_cell_text(flow.cell(row_index, 1), block, fill, 12)
        set_cell_text(flow.cell(row_index, 2), desc, "FFFFFF", 11, False)

    add_heading(doc, "3. Chú Giải Màu")
    legend = doc.add_table(rows=6, cols=2)
    legend.style = "Table Grid"
    style_table(legend)

    legend_rows = [
        ("Màu vàng", "Người dùng / giáo viên", "FEF3C7"),
        ("Màu xanh dương", "Giao diện web: form, preview, toolbar", "EFF6FF"),
        ("Màu tím", "Backend API: validate, generate, refine, export", "F3E8FF"),
        ("Màu xanh lá", "AI/OCR providers: Gemini Vision, GPT models", "ECFDF5"),
        ("Màu cam", "Đầu ra: DOCX, PDF, copy nội dung", "FFF7ED"),
        ("Màu xám", "Lưu nháp, autosave, version history", "F9FAFB"),
    ]

    for row_index, (color_name, meaning, fill) in enumerate(legend_rows):
        set_cell_text(legend.cell(row_index, 0), color_name, fill, 11)
        set_cell_text(legend.cell(row_index, 1), meaning, "FFFFFF", 11, False)

    add_paragraph(
        doc,
        "Kết luận: EduPlan AI gồm 5 lớp chính: nhập liệu, xử lý AI/OCR, kiểm tra chất lượng, hiển thị preview đẹp và xuất file Word/PDF.",
        12,
        True,
    )

    doc.save(OUT_FILE)
    print(OUT_FILE)


if __name__ == "__main__":
    main()
